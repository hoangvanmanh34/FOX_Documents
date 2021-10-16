import os
import docx

class ChoiChoi:
    def __init__(self, *args):
        print('hello')
        self.istrCheck = 'BBPPBB'
        self.iB = 0
        self.iP = 0
        self.load_content()

    def load_content(self):
        self.iB = 0
        self.iP = 0
        icontent = ''
        docfile = docx.Document(r'FILE_20211016_134230_data_bcr.docx')#('hehe.docx')#(r'FILE_20211016_134230_data_bcr.docx')
        iparags = docfile.paragraphs
        for para in iparags:
            if str(para.text).strip() != '':
                #print(para.text)
                iliststr = str(para.text).strip().split(self.istrCheck)
                #print(iliststr)
                if str(para.text).strip().find(self.istrCheck) == 0:
                    if iliststr[0].strip() != '':
                        self.Check_item(iliststr[0])
                if len(iliststr) > 1:
                    for i in range(1, len(iliststr)):
                        if iliststr[i].strip() != '':
                            self.Check_item(iliststr[i])
        itotal = (self.iB+self.iP)
        print('Number of B: ', self.iB)
        print('Number of P: ', self.iP)

        print('Percent of B: ', round(self.iB/itotal, 2)*100, '(%)')
        print('Percent of P: ', round(self.iP/itotal, 2)*100, '(%)')

    def load_content2(self):
        self.iB = 0
        self.iP = 0
        icontent = ''
        docfile = docx.Document(r'FILE_20211016_134230_data_bcr.docx')#('hehe.docx')#(r'FILE_20211016_134230_data_bcr.docx')
        iparags = docfile.paragraphs
        for para in iparags:
            if str(para.text).strip() != '':
                if str(para.text).strip().find(self.istrCheck) >= 0:                    
                        self.Check_item2(iliststr[0])
                if len(iliststr) > 1:
                    for i in range(1, len(iliststr)):
                        if iliststr[i].strip() != '':
                            self.Check_item(iliststr[i])
        itotal = (self.iB+self.iP)
        print('Number of B: ', self.iB)
        print('Number of P: ', self.iP)

        print('Percent of B: ', round(self.iB/itotal, 2)*100, '(%)')
        print('Percent of P: ', round(self.iP/itotal, 2)*100, '(%)')
                

    def Check_item(self, istr):
        if istr[0] == 'B':
            self.iB += 1
        elif istr[0] == 'P':
            self.iP += 1

    def Check_item2(self, istr):
        if istr[0] == 'B':
            self.iB += 1
            self.istrCheck = self.istrCheck + 'B'
        elif istr[0] == 'P':
            self.iP += 1
            self.istrCheck = self.istrCheck + 'P'

ChoiChoi()
